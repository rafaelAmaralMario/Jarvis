"""Document generation — DOCX, PDF, XLSX."""

import logging
import os
from datetime import datetime

logger = logging.getLogger(__name__)


class DocumentGenerator:
    def create_docx(self, path: str, title: str = "", sections: list | None = None) -> dict:
        try:
            from docx import Document
            from docx.shared import Inches, Pt
            from docx.enum.text import WD_ALIGN_PARAGRAPH
        except ImportError:
            return {"success": False, "error": "python-docx not installed"}

        doc = Document()
        if title:
            doc.add_heading(title, level=0)

        for sec in (sections or []):
            heading = sec.get("heading", "")
            if heading:
                doc.add_heading(heading, level=sec.get("level", 1))

            content = sec.get("content", "")
            stype = sec.get("type", "paragraph")

            if stype == "paragraph":
                p = doc.add_paragraph(content)
                if sec.get("bold"):
                    for run in p.runs:
                        run.bold = True
                if sec.get("italic"):
                    for run in p.runs:
                        run.italic = True
            elif stype == "table":
                rows = sec.get("rows", [])
                if rows:
                    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
                    table.style = "Table Grid"
                    for i, row_data in enumerate(rows):
                        for j, cell_val in enumerate(row_data):
                            table.cell(i, j).text = str(cell_val)
            elif stype == "list":
                for item in content if isinstance(content, list) else [content]:
                    doc.add_paragraph(item, style="List Bullet")

        os.makedirs(os.path.dirname(path) or ".", exist_ok=True)
        doc.save(path)
        size = os.path.getsize(path)
        return {"success": True, "path": path, "size_bytes": size}

    def create_pdf(self, path: str, sections: list | None = None) -> dict:
        try:
            from reportlab.lib.pagesizes import A4
            from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
            from reportlab.lib.units import mm
            from reportlab.platypus import SimpleDocTemplate, Paragraph, Table, TableStyle, Spacer
            from reportlab.lib import colors
        except ImportError:
            return {"success": False, "error": "reportlab not installed"}

        doc = SimpleDocTemplate(path, pagesize=A4)
        styles = getSampleStyleSheet()
        flowables = []

        for sec in (sections or []):
            heading = sec.get("heading", "")
            if heading:
                level = sec.get("level", 1)
                style_name = f"Heading{level}" if level <= 3 else "Heading3"
                flowables.append(Paragraph(heading, styles[style_name]))
                flowables.append(Spacer(1, 6))

            content = sec.get("content", "")
            stype = sec.get("type", "paragraph")

            if stype == "paragraph":
                flowables.append(Paragraph(content, styles["Normal"]))
            elif stype == "table":
                rows = sec.get("rows", [])
                if rows:
                    t = Table(rows)
                    t.setStyle(TableStyle([
                        ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
                        ("BACKGROUND", (0, 0), (-1, 0), colors.lightgrey),
                        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                    ]))
                    flowables.append(t)
            elif stype == "list":
                for item in content if isinstance(content, list) else [content]:
                    flowables.append(Paragraph(f"• {item}", styles["Normal"]))

            flowables.append(Spacer(1, 6))

        doc.build(flowables)
        size = os.path.getsize(path)
        return {"success": True, "path": path, "size_bytes": size}

    def create_xlsx(self, path: str, sheets: list | None = None) -> dict:
        try:
            from openpyxl import Workbook
            from openpyxl.styles import Font, Alignment, numbers
        except ImportError:
            return {"success": False, "error": "openpyxl not installed"}

        wb = Workbook()
        wb.remove(wb.active)

        for sheet_data in (sheets or []):
            ws = wb.create_sheet(title=sheet_data.get("name", "Sheet"))
            headers = sheet_data.get("headers", [])
            rows = sheet_data.get("rows", [])
            formulas = sheet_data.get("formulas", {})

            if headers:
                for col, h in enumerate(headers, 1):
                    cell = ws.cell(row=1, column=col, value=h)
                    cell.font = Font(bold=True)
                start_row = 2
            else:
                start_row = 1

            for i, row_data in enumerate(rows, start_row):
                for j, val in enumerate(row_data):
                    cell = ws.cell(row=i, column=j + 1, value=val)

            for cell_ref, formula in formulas.items():
                ws[cell_ref] = formula

        os.makedirs(os.path.dirname(path) or ".", exist_ok=True)
        wb.save(path)
        size = os.path.getsize(path)
        return {"success": True, "path": path, "size_bytes": size}
