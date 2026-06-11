"""Document reader — extract text, tables, metadata from PDF, DOCX, XLSX."""

import logging
import os
from dataclasses import dataclass, field

logger = logging.getLogger(__name__)


@dataclass
class DocumentContent:
    text: str = ""
    tables: list[list[list]] = field(default_factory=list)
    metadata: dict = field(default_factory=dict)
    pages: int = 0
    file_path: str = ""
    file_type: str = ""
    size_bytes: int = 0


class DocumentReader:
    def read_pdf(self, path: str, start_page: int = 1, end_page: int = 0) -> DocumentContent:
        import fitz
        result = DocumentContent(file_path=path, file_type="pdf")
        result.size_bytes = os.path.getsize(path)
        doc = fitz.open(path)
        result.metadata = {k: v for k, v in doc.metadata.items() if v}
        result.pages = len(doc)
        all_text = []
        all_tables = []
        page_range = range(start_page - 1, min(end_page, result.pages) if end_page else result.pages)
        for i in page_range:
            page = doc[i]
            text = page.get_text("text").strip()
            if not text:
                text = f"[Page {i + 1}: no extractable text (possible scanned image)]"
            all_text.append(f"--- Page {i + 1} ---\n{text}")
            try:
                tabs = page.find_tables()
                for t in tabs:
                    rows = [[cell.text.strip() for cell in row.cells] for row in t.extract()]
                    all_tables.append(rows)
            except Exception:
                pass
        result.text = "\n\n".join(all_text)
        result.tables = all_tables
        doc.close()
        return result

    def read_docx(self, path: str) -> DocumentContent:
        from docx import Document as DocxDocument
        from docx.shared import Pt
        result = DocumentContent(file_path=path, file_type="docx")
        result.size_bytes = os.path.getsize(path)
        doc = DocxDocument(path)
        cp = doc.core_properties
        result.metadata = {
            "author": cp.author or "",
            "title": cp.title or "",
            "created": str(cp.created) if cp.created else "",
            "modified": str(cp.modified) if cp.modified else "",
        }
        paragraphs = []
        for p in doc.paragraphs:
            if p.text.strip():
                paragraphs.append(p.text)
        result.text = "\n".join(paragraphs)
        tables = []
        for table in doc.tables:
            rows = []
            for row in table.rows:
                rows.append([cell.text.strip() for cell in row.cells])
            tables.append(rows)
        result.tables = tables
        result.pages = len(paragraphs)
        return result

    def read_xlsx(self, path: str, sheet: str = "") -> DocumentContent:
        from openpyxl import load_workbook
        result = DocumentContent(file_path=path, file_type="xlsx")
        result.size_bytes = os.path.getsize(path)
        wb = load_workbook(path, data_only=True)
        cp = wb.properties
        result.metadata = {
            "creator": cp.creator or "",
            "created": str(cp.created) if cp.created else "",
            "modified": str(cp.modified) if cp.modified else "",
            "sheets": wb.sheetnames,
        }
        wb_sheet = wb[sheet] if sheet else wb.active
        result.pages = len(wb.sheetnames)
        rows_data = []
        for row in wb_sheet.iter_rows(values_only=True):
            rows_data.append([str(c) if c is not None else "" for c in row])
        result.tables = [rows_data]
        result.text = "\n".join(
            " | ".join(str(c) if c is not None else "" for c in row)
            for row in wb_sheet.iter_rows(values_only=True)
        )
        wb.close()
        return result
